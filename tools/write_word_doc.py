"""
WAT tool: Generate a Word (.docx) document from a JSON content spec.

Input JSON schema (--input or --input-file):
{
  "output_path": "path/to/output.docx",
  "font_name": "Times New Roman",          // optional, default Times New Roman
  "heading_font_size": 14,                 // optional, default 14
  "body_font_size": 12,                    // optional, default 12
  "title": "Optional document title (Heading 1)",
  "sections": [
    {
      "image": {                            // optional — insert an image
        "path": "/abs/path/to/image.png",
        "width_inches": 2.0,
        "align": "center"                  // "left" | "center" | "right"
      },
      "heading": "Section heading text",   // optional
      "heading_level": 2,                  // 1–4, default 2
      "paragraphs": ["Para 1", "..."],     // optional
      "paragraph_align": "left",           // "left" | "center" | "right"
      "bullet_list": ["Item 1", "..."],    // optional
      "table": {                           // optional
        "headers": ["Col A", "Col B"],
        "rows": [["val1", "val2"]]
      },
      "flowchart": [                        // optional — vertical flow diagram
        {"box": "text", "color": "RRGGBB", "text_color": "RRGGBB", "bold": true},
        {"arrow": "↓"}
      ]
    }
  ]
}
"""

import argparse
import json
import os
import sys

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from dotenv import load_dotenv

load_dotenv()

# ──────────────────────────────────────────────
# XML / style helpers
# ──────────────────────────────────────────────

def _apply_font(run, font_name: str, font_size: int, bold: bool = False, color_hex: str = None):
    run.font.name = font_name
    run.font.size = Pt(font_size)
    if bold:
        run.bold = True
    if color_hex:
        h = color_hex.lstrip("#")
        run.font.color.rgb = RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _set_doc_styles(doc, font_name: str, body_size: int, heading_size: int):
    """Apply font, sizes, and black color to Normal and Heading 1–4 styles."""
    normal = doc.styles["Normal"]
    normal.font.name = font_name
    normal.font.size = Pt(body_size)
    normal.font.color.rgb = RGBColor(0, 0, 0)

    sizes = {1: heading_size + 2, 2: heading_size, 3: heading_size, 4: body_size}
    for lvl in range(1, 5):
        try:
            s = doc.styles[f"Heading {lvl}"]
            s.font.name = font_name
            s.font.size = Pt(sizes[lvl])
            s.font.color.rgb = RGBColor(0, 0, 0)  # override Word's default blue theme color
        except KeyError:
            pass


def _shade_cell(cell, fill_hex: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex.lstrip("#"))
    tcPr.append(shd)


def _add_cell_border(cell, color_hex: str = "888888"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), color_hex.lstrip("#"))
        borders.append(b)
    tcPr.append(borders)


def _center_table(table):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "center")
    tblPr.append(jc)


# ──────────────────────────────────────────────
# Section renderers
# ──────────────────────────────────────────────

def _render_image(doc, image_spec: dict):
    path = image_spec.get("path", "")
    if not os.path.exists(path):
        doc.add_paragraph(f"[Image not found: {path}]")
        return
    width = image_spec.get("width_inches", 2.0)
    align_str = image_spec.get("align", "center")
    align_map = {"center": WD_ALIGN_PARAGRAPH.CENTER, "right": WD_ALIGN_PARAGRAPH.RIGHT,
                 "left": WD_ALIGN_PARAGRAPH.LEFT}
    p = doc.add_paragraph()
    p.alignment = align_map.get(align_str, WD_ALIGN_PARAGRAPH.CENTER)
    p.add_run().add_picture(path, width=Inches(width))


def _render_flowchart(doc, items: list, font_name: str, font_size: int):
    """
    Vertical flow diagram as a borderless 1-column table.
    Box items get colored backgrounds and explicit borders.
    Arrow items are transparent rows with centered ↓ text.
    """
    table = doc.add_table(rows=len(items), cols=1)
    _center_table(table)

    for i, item in enumerate(items):
        cell = table.rows[i].cells[0]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if "arrow" in item:
            run = p.add_run(item["arrow"])
            _apply_font(run, font_name, font_size + 4, bold=True)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
        else:
            text = item.get("box", "")
            fill = item.get("color", "E8E8E8")
            text_color = item.get("text_color")
            bold = item.get("bold", False)

            p.paragraph_format.space_before = Pt(5)
            p.paragraph_format.space_after = Pt(5)
            run = p.add_run(text)
            _apply_font(run, font_name, font_size, bold=bold, color_hex=text_color)
            _shade_cell(cell, fill)
            _add_cell_border(cell)

    doc.add_paragraph("")


def _render_table(doc, tbl_spec: dict, font_name: str, body_size: int):
    headers = tbl_spec.get("headers", [])
    rows = tbl_spec.get("rows", [])
    col_count = len(headers) or (len(rows[0]) if rows else 1)
    table = doc.add_table(rows=1 + len(rows), cols=col_count)
    table.style = "Table Grid"

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(str(h))
        _apply_font(run, font_name, body_size, bold=True)

    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            _apply_font(run, font_name, body_size)


# ──────────────────────────────────────────────
# Main builder
# ──────────────────────────────────────────────

def build_document(spec: dict) -> str:
    output_path = spec.get("output_path")
    if not output_path:
        raise ValueError("'output_path' is required in the input spec")

    font_name = spec.get("font_name", "Times New Roman")
    heading_size = spec.get("heading_font_size", 14)
    body_size = spec.get("body_font_size", 12)

    doc = Document()
    _set_doc_styles(doc, font_name, body_size, heading_size)

    for sec in doc.sections:
        sec.top_margin = Inches(1.0)
        sec.bottom_margin = Inches(1.0)
        sec.left_margin = Inches(1.2)
        sec.right_margin = Inches(1.2)

    if spec.get("title"):
        doc.add_heading(spec["title"], level=1)

    align_map = {
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "left": WD_ALIGN_PARAGRAPH.LEFT,
    }

    for section in spec.get("sections", []):
        if section.get("image"):
            _render_image(doc, section["image"])

        if section.get("heading"):
            doc.add_heading(section["heading"], level=int(section.get("heading_level", 2)))

        p_align = align_map.get(section.get("paragraph_align", "left"), WD_ALIGN_PARAGRAPH.LEFT)
        for para in section.get("paragraphs", []):
            p = doc.add_paragraph()
            p.alignment = p_align
            _apply_font(p.add_run(para), font_name, body_size)

        for item in section.get("bullet_list", []):
            p = doc.add_paragraph(style="List Bullet")
            _apply_font(p.add_run(item), font_name, body_size)

        if section.get("table"):
            _render_table(doc, section["table"], font_name, body_size)

        if section.get("flowchart"):
            _render_flowchart(doc, section["flowchart"], font_name, body_size)

        doc.add_paragraph("")

    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    doc.save(output_path)
    return output_path


# ──────────────────────────────────────────────
# CLI entry point
# ──────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description="Write a Word .docx file from a JSON spec")
    group = parser.add_mutually_exclusive_group(required=True)
    group.add_argument("--input", help="JSON string spec")
    group.add_argument("--input-file", help="Path to JSON spec file")
    parser.add_argument("--output-file", help="Optional path to write result JSON")
    args = parser.parse_args()

    try:
        if args.input_file:
            with open(args.input_file) as f:
                spec = json.load(f)
        else:
            spec = json.loads(args.input)

        saved_path = build_document(spec)
        result = {"status": "ok", "output_path": saved_path}
    except Exception as e:
        print(json.dumps({"status": "error", "message": str(e)}), file=sys.stderr)
        sys.exit(1)

    output = json.dumps(result, indent=2)
    print(output)

    if args.output_file:
        os.makedirs(os.path.dirname(args.output_file) or ".", exist_ok=True)
        with open(args.output_file, "w") as f:
            f.write(output)


if __name__ == "__main__":
    main()
